#!/usr/bin/env python3
"""Build the ATS-friendly Verily Platform resume.

Design system:
- Base preset: compact_reference_guide.
- Header pattern: customer_pack, adapted as a left-aligned resume masthead.
- Named resume overrides: Letter page, 0.62" side margins, 0.52" top/bottom
  margins, Arial 9.4 pt body, compact paragraph rhythm, and no tables.
  These overrides keep the application to two pages while preserving readable
  typography and a conventional single-column ATS reading order.
"""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("resumes/generated/resume_verily.docx")

FONT = "Arial"
NAVY = RGBColor(18, 50, 74)
TEAL = RGBColor(15, 107, 109)
INK = RGBColor(27, 36, 44)
MUTED = RGBColor(83, 96, 106)
RULE = "C8D2D8"


def set_run(run, *, size=9.4, bold=False, italic=False, color=INK):
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return run


def set_keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_next = OxmlElement("w:keepNext")
    p_pr.append(keep_next)


def set_keep_lines(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    keep_lines = OxmlElement("w:keepLines")
    p_pr.append(keep_lines)


def add_bottom_border(paragraph, color=RULE, size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), fill)
    p_pr.append(shading)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.52)
    section.bottom_margin = Inches(0.52)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.28)
    section.footer_distance = Inches(0.28)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing = 1.04
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    for style_name, size, color, before, after in [
        ("Heading 1", 11.4, NAVY, 7, 3),
        ("Heading 2", 10.4, NAVY, 4, 1.5),
        ("Heading 3", 9.6, TEAL, 3, 1),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    bullet = doc.styles["List Bullet"]
    bullet.font.name = FONT
    bullet._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    bullet._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    bullet.font.size = Pt(9.2)
    bullet.font.color.rgb = INK
    bullet.paragraph_format.left_indent = Inches(0.31)
    bullet.paragraph_format.first_line_indent = Inches(-0.17)
    bullet.paragraph_format.space_before = Pt(0)
    bullet.paragraph_format.space_after = Pt(1.7)
    bullet.paragraph_format.line_spacing = 1.02

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    set_run(
        p.add_run("Dheeraj Bharat Sethi  |  Verily Platform Resume"),
        size=7.5,
        color=MUTED,
    )


def add_masthead(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("DHEERAJ BHARAT SETHI"), size=21.5, bold=True, color=NAVY)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    set_run(
        p.add_run("Senior Software Engineer  |  Platform, Distributed Systems & Production AI"),
        size=10.7,
        bold=True,
        color=TEAL,
    )

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(p.add_run("Toronto, Ontario  •  "), size=8.7, color=MUTED)
    set_run(p.add_run("dj.sethi09@gmail.com"), size=8.7, color=NAVY)
    set_run(p.add_run("  •  linkedin.com/in/dheerajbharatsethi"), size=8.7, color=NAVY)
    set_run(p.add_run("  •  github.com/dheerix"), size=8.7, color=NAVY)
    add_bottom_border(p, color="7DA8AA", size="10", space="5")


def add_section_heading(doc, text):
    p = doc.add_paragraph(text, style="Heading 1")
    set_keep_with_next(p)
    add_bottom_border(p, color=RULE, size="5", space="2")
    return p


def add_role_heading(doc, company, role, dates=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3.5)
    p.paragraph_format.space_after = Pt(0.8)
    set_keep_with_next(p)
    set_run(p.add_run(company), size=10.2, bold=True, color=NAVY)
    set_run(p.add_run(f"  |  {role}"), size=9.7, bold=True, color=INK)
    if dates:
        set_run(p.add_run(f"  |  {dates}"), size=8.8, color=MUTED)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    set_run(p.add_run(text), size=9.2, color=INK)
    set_keep_lines(p)
    return p


def add_label_line(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2.2)
    set_run(p.add_run(f"{label}: "), size=9.1, bold=True, color=NAVY)
    set_run(p.add_run(value), size=9.1, color=INK)
    return p


def build():
    doc = Document()
    configure_document(doc)
    add_masthead(doc)

    add_section_heading(doc, "PROFESSIONAL SUMMARY")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    set_run(
        p.add_run(
            "Senior software engineer with 14+ years designing and delivering distributed "
            "software systems across cloud platforms, production AI, backend services, and "
            "modern web applications. Experienced leading end-to-end delivery from architecture "
            "through production operations, translating ambiguous product needs into reliable, "
            "observable, and maintainable systems. Brings automotive marketplace experience and "
            "hands-on healthcare engineering within HIPAA-governed environments."
        ),
        size=9.5,
    )

    add_section_heading(doc, "CORE STRENGTHS")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    shade_paragraph(p, "F1F6F6")
    set_run(
        p.add_run(
            "Distributed Systems  •  Platform Engineering  •  Technical Leadership  •  "
            "Full-Stack Delivery  •  Cloud Architecture  •  Production AI  •  APIs  •  "
            "Reliability & Observability  •  Incremental Modernization  •  Technical Documentation"
        ),
        size=8.9,
        bold=True,
        color=NAVY,
    )

    add_section_heading(doc, "PROFESSIONAL EXPERIENCE")
    add_role_heading(
        doc,
        "OPENLANE",
        "Senior Software Engineer — Marketplace, AI & Platform Engineering",
        "2022–Present",
    )
    add_bullet(
        doc,
        "Led incremental modernization of a legacy Oracle- and Java-based automotive marketplace "
        "through cloud-native services and event-driven integrations, enabling new capabilities "
        "while preserving production continuity.",
    )
    add_bullet(
        doc,
        "Owned modernization of the Vehicle Detail experience and its backend-for-frontend, "
        "aggregating multiple internal services behind a cohesive API and customer-facing workflow.",
    )
    add_bullet(
        doc,
        "Designed and delivered React, Svelte, and Stencil micro frontends and contributed reusable "
        "UI patterns that improved consistency across marketplace product surfaces.",
    )
    add_bullet(
        doc,
        "Helped evolve Guardlane from a working AI moderation flow toward a production-ready system "
        "combining ML classification, model evaluation, LLM fallback, human review, governance, "
        "observability, and operational controls.",
    )
    add_bullet(
        doc,
        "Contributed to AI Upload, integrating validation, AI-assisted content generation, review "
        "paths, backend services, and frontend workflows into a real dealer operating process.",
    )
    add_bullet(
        doc,
        "Migrated services from AWS SDK for JavaScript v2 to v3, addressing client initialization, "
        "packaging, runtime behavior, error handling, testing, and deployment confidence.",
    )
    add_bullet(
        doc,
        "Supported production operations and incident investigations using distributed tracing, "
        "logs, metrics, Honeycomb, and OpenTelemetry; converted findings into durable reliability improvements.",
    )
    add_bullet(
        doc,
        "Drove technical discussions, onboarding, design documentation, and cross-functional alignment "
        "with product, engineering, data, and operations partners; kept delivery moving during periods "
        "when the technical lead was unavailable.",
    )

    doc.add_page_break()

    add_role_heading(
        doc,
        "HOLLAND & BARRETT",
        "Software Engineer — Enterprise Search & Commerce",
        "2018–2022",
    )
    add_bullet(
        doc,
        "Built backend services and search APIs supporting product discovery across a large, changing "
        "commerce catalog, integrating Elasticsearch with customer-facing workflows.",
    )
    add_bullet(
        doc,
        "Delivered React and JavaScript interfaces for search, promotions, and price-definition workflows, "
        "working across frontend and backend boundaries as delivery required.",
    )
    add_bullet(
        doc,
        "Enhanced ANTLR-based business rules and promotion logic that connected merchandising intent "
        "to search and commerce behavior.",
    )
    add_bullet(
        doc,
        "Investigated production issues and collaborated with product managers and globally distributed "
        "engineers to modernize business-critical search capabilities.",
    )

    add_role_heading(doc, "FOODMESH", "Software Engineer — Sustainability Platform")
    add_bullet(
        doc,
        "Built backend, frontend, and mobile capabilities using Python, Django, React, and Flutter for "
        "customer-facing and operational food-recovery workflows.",
    )
    add_bullet(
        doc,
        "Translated lightly specified product needs into maintainable features while collaborating "
        "directly with operational stakeholders and distributed team members.",
    )

    add_role_heading(doc, "RHEUMERA", "Software Engineer — Healthcare Platform")
    add_bullet(
        doc,
        "Developed workflow-oriented healthcare software using React and Spring Boot in an environment "
        "requiring reliable, privacy-conscious, HIPAA-aligned engineering practices.",
    )
    add_bullet(
        doc,
        "Contributed maintainable product capabilities where data handling, correctness, access control, "
        "and clear operational behavior were essential.",
    )

    add_section_heading(doc, "TECHNICAL SKILLS")
    add_label_line(
        doc,
        "Cloud & Platform",
        "AWS, Google Cloud Platform (GCP), Lambda, API Gateway, S3, SNS, SQS, Kinesis, "
        "Docker, Kubernetes, Terraform, ArgoCD, Azure DevOps, CI/CD",
    )
    add_label_line(
        doc,
        "Backend & Architecture",
        "Java, C#, .NET, Node.js, Python, Spring Boot, REST, GraphQL, microservices, "
        "distributed systems, event-driven architecture, backend-for-frontend",
    )
    add_label_line(
        doc,
        "Frontend",
        "React, TypeScript, JavaScript, Svelte, Stencil, Flutter, HTML, CSS, responsive web development",
    )
    add_label_line(
        doc,
        "Data & Messaging",
        "PostgreSQL, Oracle, SQL Server, DynamoDB, MongoDB, Elasticsearch, Kafka, RabbitMQ, Pulsar",
    )
    add_label_line(
        doc,
        "AI & Operations",
        "SageMaker, ML workflows, model training and evaluation, LLM integration, human-in-the-loop "
        "systems, MLOps, AI governance, Honeycomb, OpenTelemetry, logging, metrics, tracing",
    )
    add_label_line(
        doc,
        "Security & Healthcare",
        "HIPAA-aware engineering, authentication, authorization, secure APIs, privacy-conscious data handling",
    )

    add_section_heading(doc, "EDUCATION")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    set_run(p.add_run("University of Windsor"), size=10, bold=True, color=NAVY)
    set_run(
        p.add_run("  |  Master’s Degree, Electrical and Computer Engineering  |  2022–2024"),
        size=9.3,
        bold=True,
        color=INK,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run(
        p.add_run(
            "Selected project: designed an image-caption generation system using neural networks and "
            "LSTM sequence modeling, trained and evaluated on the Flickr8k dataset."
        ),
        size=8.9,
        color=MUTED,
    )

    # Word compatibility and metadata.
    doc.core_properties.title = "Dheeraj Bharat Sethi — Verily Platform Resume"
    doc.core_properties.subject = (
        "Senior Software Engineer application for Verily Health platform, technical lead, "
        "precision health, full-stack, and developer platform roles"
    )
    doc.core_properties.author = "Dheeraj Bharat Sethi"
    doc.core_properties.keywords = (
        "platform engineering, distributed systems, technical leadership, React, TypeScript, "
        "AWS, GCP, Terraform, Kubernetes, production AI, HIPAA, Verily"
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
